"""
文件解析模块 - 支持多种文档格式
"""
import os
import magic
import pdfplumber
from docx import Document
import extract_msg
from typing import Optional


class FileParser:
    """文件解析器，支持PDF、Word、Email等格式"""

    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """
        自动识别文件类型
        :param file_path: 文件路径
        :return: 文件MIME类型
        """
        mime = magic.Magic(mime=True)
        return mime.from_file(file_path)

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """
        解析PDF文件
        :param file_path: PDF文件路径
        :return: 提取的文本内容
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")
        return text.strip()

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """
        解析DOCX文件
        :param file_path: DOCX文件路径
        :return: 提取的文本内容
        """
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text += "\n" + row_text

            return text.strip()
        except Exception as e:
            raise ValueError(f"DOCX解析失败: {str(e)}")

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """
        解析TXT文件
        :param file_path: TXT文件路径
        :return: 文本内容
        """
        try:
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read().strip()
                except UnicodeDecodeError:
                    continue
            raise ValueError("无法识别文件编码")
        except Exception as e:
            raise ValueError(f"TXT解析失败: {str(e)}")

    @staticmethod
    def parse_msg(file_path: str) -> str:
        """
        解析MSG文件（Outlook邮件）
        :param file_path: MSG文件路径
        :return: 邮件内容
        """
        try:
            msg = extract_msg.Message(file_path)

            # 提取邮件基本信息
            text = f"发件人: {msg.sender}\n"
            text += f"收件人: {msg.to}\n"
            text += f"主题: {msg.subject}\n"
            text += f"日期: {msg.date}\n"
            text += f"\n邮件正文:\n{msg.body}\n"

            msg.close()
            return text.strip()
        except Exception as e:
            raise ValueError(f"MSG解析失败: {str(e)}")

    @staticmethod
    def parse_eml(file_path: str) -> str:
        """
        解析EML文件（标准邮件格式）
        :param file_path: EML文件路径
        :return: 邮件内容
        """
        import email
        from email import policy

        try:
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f, policy=policy.default)

            # 提取邮件信息
            text = f"发件人: {msg.get('From')}\n"
            text += f"收件人: {msg.get('To')}\n"
            text += f"主题: {msg.get('Subject')}\n"
            text += f"日期: {msg.get('Date')}\n"

            # 提取正文
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text += f"\n邮件正文:\n{part.get_content()}\n"
                        break
            else:
                text += f"\n邮件正文:\n{msg.get_content()}\n"

            return text.strip()
        except Exception as e:
            raise ValueError(f"EML解析失败: {str(e)}")

    def parse_file(self, file_path: str) -> str:
        """
        根据文件类型自动选择解析器
        :param file_path: 文件路径
        :return: 提取的文本内容
        """
        # 检测文件类型
        mime_type = self.detect_file_type(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()

        # 根据MIME类型或扩展名选择解析器
        if 'pdf' in mime_type or file_ext == '.pdf':
            return self.parse_pdf(file_path)

        elif 'word' in mime_type or 'document' in mime_type or file_ext in ['.docx', '.doc']:
            if file_ext == '.docx':
                return self.parse_docx(file_path)
            else:
                # .doc文件暂时不支持，建议用户转换为.docx
                raise ValueError("暂不支持.doc格式，请转换为.docx后上传")

        elif file_ext == '.msg':
            return self.parse_msg(file_path)

        elif file_ext == '.eml':
            return self.parse_eml(file_path)

        elif 'text' in mime_type or file_ext == '.txt':
            return self.parse_txt(file_path)

        else:
            raise ValueError(f"不支持的文件格式: {mime_type} ({file_ext})")


# 测试代码
if __name__ == "__main__":
    parser = FileParser()

    # 测试PDF
    # text = parser.parse_file("sample.pdf")
    # print(text)
